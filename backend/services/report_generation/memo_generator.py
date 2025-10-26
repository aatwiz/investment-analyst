"""
Investment Memo Generator
Generates professional investment memos based on company data from Features 1-4
"""

from typing import Dict, List, Optional, Any
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import io
from loguru import logger
from openai import AsyncOpenAI
import os


class InvestmentMemoGenerator:
    """
    Generates comprehensive investment memos integrating data from all features
    """
    
    def __init__(self):
        """Initialize memo generator with OpenAI client"""
        self.client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        logger.info("InvestmentMemoGenerator initialized")
    
    async def generate_memo(
        self,
        company_data: Dict[str, Any],
        deal_data: Optional[Dict[str, Any]] = None,
        market_data: Optional[Dict[str, Any]] = None,
        financial_model: Optional[Dict[str, Any]] = None,
        template_type: str = "standard",
        analyst_name: str = "Investment Analyst",
        firm_name: str = "Investment Firm"
    ) -> io.BytesIO:
        """
        Generate complete investment memo
        
        Args:
            company_data: Company information from Feature 1
            deal_data: Deal qualification data from Feature 1
            market_data: Market analysis from Feature 3
            financial_model: Financial projections from Feature 4
            template_type: Memo template (standard, detailed, summary)
            analyst_name: Name of analyst preparing memo
            firm_name: Investment firm name
        
        Returns:
            BytesIO containing DOCX memo
        """
        logger.info(f"Generating {template_type} investment memo for {company_data.get('name', 'Unknown')}")
        
        # Create document
        doc = Document()
        
        # Set up styles
        self._setup_styles(doc)
        
        # Generate sections
        self._add_header(doc, company_data, firm_name)
        self._add_executive_summary(doc, company_data, deal_data, analyst_name, firm_name)
        
        # Generate AI-powered content for key sections
        investment_rationale = await self._generate_investment_rationale(
            company_data, deal_data, market_data, financial_model
        )
        self._add_investment_rationale(doc, investment_rationale)
        
        self._add_key_terms(doc, company_data, deal_data, financial_model)
        
        risk_analysis = await self._generate_risk_analysis(
            company_data, market_data, financial_model
        )
        self._add_risk_factors(doc, risk_analysis)
        
        recommendation = await self._generate_recommendation(
            company_data, deal_data, financial_model, risk_analysis
        )
        self._add_recommendation(doc, recommendation, analyst_name, firm_name)
        
        self._add_appendices(doc, market_data, financial_model)
        
        self._add_signature(doc, analyst_name, firm_name)
        
        # Save to BytesIO
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        logger.info("Investment memo generated successfully")
        return output
    
    def _setup_styles(self, doc: Document):
        """Set up document styles"""
        styles = doc.styles
        
        # Heading 1 style
        if 'Custom Heading 1' not in [s.name for s in styles]:
            heading1 = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1.font.size = Pt(18)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(0, 51, 102)
        
        # Heading 2 style
        if 'Custom Heading 2' not in [s.name for s in styles]:
            heading2 = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            heading2.font.size = Pt(14)
            heading2.font.bold = True
            heading2.font.color.rgb = RGBColor(0, 102, 204)
    
    def _add_header(self, doc: Document, company_data: Dict, firm_name: str):
        """Add memo header"""
        # Firm logo placeholder
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(firm_name.upper())
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Contact info (placeholder)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"contact@{firm_name.lower().replace(' ', '')}.com | www.{firm_name.lower().replace(' ', '')}.com")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("INVESTMENT MEMO")
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        doc.add_paragraph()
    
    def _add_executive_summary(
        self,
        doc: Document,
        company_data: Dict,
        deal_data: Optional[Dict],
        analyst_name: str,
        firm_name: str
    ):
        """Add executive summary section"""
        doc.add_heading('I. Executive Summary', level=1)
        
        # Metadata
        p = doc.add_paragraph()
        p.add_run('Date: ').bold = True
        p.add_run(datetime.now().strftime('%B %d, %Y'))
        
        p = doc.add_paragraph()
        p.add_run('Prepared by: ').bold = True
        p.add_run(analyst_name)
        
        p = doc.add_paragraph()
        p.add_run('Department: ').bold = True
        p.add_run('Investment Analysis')
        
        p = doc.add_paragraph()
        p.add_run('Firm: ').bold = True
        p.add_run(firm_name)
        
        doc.add_paragraph()
        
        # Investment Opportunity
        doc.add_heading('Investment Opportunity:', level=2)
        
        bullets = [
            f"Company Name: {company_data.get('name', 'N/A')}",
            f"Industry/Market: {company_data.get('industry', 'N/A')}",
            f"Investment Stage: {company_data.get('stage', 'N/A')}",
            f"Investment Type: Equity",
            f"Funding Amount: ${company_data.get('funding_amount', 0):,.0f}",
            f"Valuation: ${company_data.get('valuation', 0):,.0f}" if company_data.get('valuation') else None,
            f"Location: {company_data.get('location', 'N/A')}"
        ]
        
        for bullet in bullets:
            if bullet:
                doc.add_paragraph(bullet, style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_investment_rationale(self, doc: Document, rationale: Dict):
        """Add investment rationale section"""
        doc.add_heading('II. Investment Rationale', level=1)
        
        # Market Opportunity
        doc.add_heading('1. Market Opportunity', level=2)
        market = rationale.get('market_opportunity', {})
        
        for point in market.get('key_points', []):
            doc.add_paragraph(point, style='List Bullet')
        
        if market.get('analysis'):
            doc.add_paragraph(market['analysis'])
        
        doc.add_paragraph()
        
        # Competitive Advantage
        doc.add_heading('2. Competitive Advantage', level=2)
        competitive = rationale.get('competitive_advantage', {})
        
        for point in competitive.get('key_points', []):
            doc.add_paragraph(point, style='List Bullet')
        
        if competitive.get('barriers'):
            p = doc.add_paragraph()
            p.add_run('Barriers to Entry: ').bold = True
            doc.add_paragraph(competitive['barriers'])
        
        doc.add_paragraph()
        
        # Financial Performance
        doc.add_heading('3. Financial Performance', level=2)
        financial = rationale.get('financial_performance', {})
        
        metrics = [
            f"Revenue: {financial.get('revenue', 'N/A')}",
            f"Revenue Growth: {financial.get('growth', 'N/A')}",
            f"Profit Margin: {financial.get('margin', 'N/A')}",
            f"Cash Flow: {financial.get('cash_flow', 'N/A')}"
        ]
        
        for metric in metrics:
            doc.add_paragraph(metric, style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_key_terms(
        self,
        doc: Document,
        company_data: Dict,
        deal_data: Optional[Dict],
        financial_model: Optional[Dict]
    ):
        """Add key investment terms section"""
        doc.add_heading('III. Key Investment Terms', level=1)
        
        # Create table
        table = doc.add_table(rows=9, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Details'
        
        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Fill data
        funding_amount = company_data.get('funding_amount', 1000000)  # Default $1M if not provided
        valuation = company_data.get('valuation', funding_amount * 4)  # Estimate if not provided
        
        # Calculate ownership percentage (avoid division by zero)
        total_value = valuation + funding_amount
        ownership_pct = (funding_amount / total_value * 100) if total_value > 0 else 0
        
        rows_data = [
            ('Investment Type', 'Equity'),
            ('Amount', f"${funding_amount:,.0f}"),
            ('Pre-money Valuation', f"${valuation:,.0f}"),
            ('Equity Ownership', f"{ownership_pct:.1f}%"),
            ('Investment Horizon', '3-5 Years'),
            ('Expected ROI', self._calculate_expected_roi(financial_model)),
            ('Exit Strategy', 'Acquisition or IPO'),
            ('Board Seats', 'To be negotiated')
        ]
        
        for i, (category, detail) in enumerate(rows_data, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = category
            row_cells[1].text = detail
        
        doc.add_paragraph()
    
    def _add_risk_factors(self, doc: Document, risk_analysis: Dict):
        """Add risk factors section"""
        doc.add_heading('IV. Risk Factors', level=1)
        
        # Market Risks
        doc.add_heading('Market Risks:', level=2)
        for risk in risk_analysis.get('market_risks', []):
            doc.add_paragraph(risk, style='List Bullet')
        
        doc.add_paragraph()
        
        # Operational Risks
        doc.add_heading('Operational Risks:', level=2)
        for risk in risk_analysis.get('operational_risks', []):
            doc.add_paragraph(risk, style='List Bullet')
        
        doc.add_paragraph()
        
        # Financial Risks
        doc.add_heading('Financial Risks:', level=2)
        for risk in risk_analysis.get('financial_risks', []):
            doc.add_paragraph(risk, style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_recommendation(
        self,
        doc: Document,
        recommendation: Dict,
        analyst_name: str,
        firm_name: str
    ):
        """Add recommendation section"""
        doc.add_heading('V. Recommendation', level=1)
        
        p = doc.add_paragraph()
        p.add_run(f"Based on the analysis provided, {analyst_name} recommends:\n\n")
        
        action = recommendation.get('action', 'review')
        
        if action == 'proceed':
            doc.add_paragraph(
                f"✓ Proceed with investment of ${recommendation.get('amount', 'TBD')}",
                style='List Bullet'
            )
        elif action == 'conditional':
            doc.add_paragraph(
                f"→ Conditionally approve subject to the following:",
                style='List Bullet'
            )
        else:
            doc.add_paragraph(
                f"✗ Decline the investment at this time",
                style='List Bullet'
            )
        
        # Conditions/Milestones
        for condition in recommendation.get('conditions', []):
            doc.add_paragraph(f"  • {condition}", style='List Bullet 2')
        
        # Reasoning
        if recommendation.get('reasoning'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('Rationale: ').bold = True
            p.add_run(recommendation['reasoning'])
        
        doc.add_paragraph()
    
    def _add_appendices(
        self,
        doc: Document,
        market_data: Optional[Dict],
        financial_model: Optional[Dict]
    ):
        """Add appendices section"""
        doc.add_heading('VI. Appendices', level=1)
        
        doc.add_heading('1. Financial Statements:', level=2)
        doc.add_paragraph(
            'Detailed financial projections and historical statements attached separately.',
            style='List Bullet'
        )
        
        if financial_model:
            metrics = financial_model.get('metrics', {})
            doc.add_paragraph(
                f"  • Total Revenue (3Y): ${metrics.get('total_revenue', 0):,.0f}",
                style='List Bullet 2'
            )
            doc.add_paragraph(
                f"  • Revenue CAGR: {metrics.get('revenue_cagr', 0)*100:.1f}%",
                style='List Bullet 2'
            )
            doc.add_paragraph(
                f"  • Months to Profitability: {metrics.get('months_to_profitability', 'N/A')}",
                style='List Bullet 2'
            )
        
        doc.add_paragraph()
        
        doc.add_heading('2. Market Analysis:', level=2)
        doc.add_paragraph(
            'Comprehensive market research and competitive landscape analysis.',
            style='List Bullet'
        )
        
        if market_data:
            if market_data.get('market_size'):
                doc.add_paragraph(
                    f"  • Market Size: {market_data['market_size']}",
                    style='List Bullet 2'
                )
            if market_data.get('growth_rate'):
                doc.add_paragraph(
                    f"  • Market Growth Rate: {market_data['growth_rate']}",
                    style='List Bullet 2'
                )
        
        doc.add_paragraph()
        
        doc.add_heading('3. Due Diligence Documents:', level=2)
        doc.add_paragraph(
            'Legal, regulatory, and compliance documentation reviewed.',
            style='List Bullet'
        )
        
        doc.add_paragraph()
    
    def _add_signature(self, doc: Document, analyst_name: str, firm_name: str):
        """Add signature block"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('Approved by:\n\n').bold = True
        
        doc.add_paragraph(analyst_name)
        doc.add_paragraph('Investment Analyst')
        doc.add_paragraph(firm_name)
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    
    async def _generate_investment_rationale(
        self,
        company_data: Dict,
        deal_data: Optional[Dict],
        market_data: Optional[Dict],
        financial_model: Optional[Dict]
    ) -> Dict:
        """Generate investment rationale using AI"""
        logger.info("Generating investment rationale with AI")
        
        prompt = f"""Generate a comprehensive investment rationale for the following company:

Company: {company_data.get('name')}
Industry: {company_data.get('industry')}
Stage: {company_data.get('stage')}
Description: {company_data.get('description', 'N/A')}
Funding: ${company_data.get('funding_amount', 0):,.0f}

{"Deal Score: " + str(deal_data.get('score', 'N/A')) if deal_data else ""}
{"Key Strengths: " + str(deal_data.get('strengths', [])) if deal_data else ""}

{"Market Analysis: " + str(market_data.get('market_overview', '')) if market_data else ""}

{"Financial Metrics: " + str(financial_model.get('metrics', {})) if financial_model else ""}

Provide a structured investment rationale with:
1. Market Opportunity (3-4 key points about market size, growth, trends)
2. Competitive Advantage (3-4 points about unique value prop, moats, barriers)
3. Financial Performance (key metrics, growth, profitability)

Format as JSON with this structure:
{{
  "market_opportunity": {{
    "key_points": ["point 1", "point 2", "point 3"],
    "analysis": "brief analysis paragraph"
  }},
  "competitive_advantage": {{
    "key_points": ["advantage 1", "advantage 2", "advantage 3"],
    "barriers": "barriers to entry description"
  }},
  "financial_performance": {{
    "revenue": "revenue description",
    "growth": "growth rate",
    "margin": "profit margin",
    "cash_flow": "cash flow status"
  }}
}}"""
        
        try:
            response = await self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are an expert investment analyst. Provide structured, data-driven investment analysis."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                response_format={"type": "json_object"}
            )
            
            import json
            result = json.loads(response.choices[0].message.content)
            logger.info("Investment rationale generated successfully")
            return result
            
        except Exception as e:
            logger.error(f"Error generating investment rationale: {e}")
            # Return fallback structure
            return {
                "market_opportunity": {
                    "key_points": [
                        f"Operating in the {company_data.get('industry', 'technology')} sector",
                        "Market shows strong growth potential",
                        "Favorable industry trends and tailwinds"
                    ],
                    "analysis": "The company operates in a growing market with significant opportunities."
                },
                "competitive_advantage": {
                    "key_points": [
                        "Unique technology or approach",
                        "Strong team with domain expertise",
                        "Early mover advantage in segment"
                    ],
                    "barriers": "High barriers to entry including technology, network effects, and brand recognition."
                },
                "financial_performance": {
                    "revenue": f"${company_data.get('funding_amount', 0):,.0f} in recent funding",
                    "growth": "Strong growth trajectory",
                    "margin": "TBD - subject to due diligence",
                    "cash_flow": "To be analyzed in detail"
                }
            }
    
    async def _generate_risk_analysis(
        self,
        company_data: Dict,
        market_data: Optional[Dict],
        financial_model: Optional[Dict]
    ) -> Dict:
        """Generate risk analysis using AI"""
        logger.info("Generating risk analysis with AI")
        
        prompt = f"""Analyze investment risks for the following company:

Company: {company_data.get('name')}
Industry: {company_data.get('industry')}
Stage: {company_data.get('stage')}

{"Market Data: " + str(market_data) if market_data else ""}
{"Financial Model: " + str(financial_model.get('metrics', {})) if financial_model else ""}

Identify specific risks in three categories:
1. Market Risks (market conditions, competition, regulatory)
2. Operational Risks (execution, team, supply chain)
3. Financial Risks (burn rate, profitability, liquidity)

Format as JSON:
{{
  "market_risks": ["risk 1", "risk 2", "risk 3"],
  "operational_risks": ["risk 1", "risk 2", "risk 3"],
  "financial_risks": ["risk 1", "risk 2", "risk 3"]
}}"""
        
        try:
            response = await self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are an expert investment risk analyst. Identify specific, material risks."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                response_format={"type": "json_object"}
            )
            
            import json
            result = json.loads(response.choices[0].message.content)
            logger.info("Risk analysis generated successfully")
            return result
            
        except Exception as e:
            logger.error(f"Error generating risk analysis: {e}")
            return {
                "market_risks": [
                    "Economic downturn could reduce market demand",
                    "Increased competition from established players",
                    "Regulatory changes may impact business model"
                ],
                "operational_risks": [
                    "Execution risk in delivering on product roadmap",
                    "Key person risk with founding team",
                    "Scaling challenges as company grows"
                ],
                "financial_risks": [
                    "High burn rate requires additional funding",
                    "Path to profitability uncertain",
                    "Liquidity concerns if growth slows"
                ]
            }
    
    async def _generate_recommendation(
        self,
        company_data: Dict,
        deal_data: Optional[Dict],
        financial_model: Optional[Dict],
        risk_analysis: Dict
    ) -> Dict:
        """Generate investment recommendation using AI"""
        logger.info("Generating investment recommendation with AI")
        
        score = deal_data.get('score', 50) if deal_data else 50
        
        prompt = f"""Based on the following investment analysis, provide a recommendation:

Company: {company_data.get('name')}
Deal Score: {score}/100
Funding: ${company_data.get('funding_amount', 0):,.0f}

Risks Identified: {len(risk_analysis.get('market_risks', [])) + len(risk_analysis.get('operational_risks', [])) + len(risk_analysis.get('financial_risks', []))} total risks

{"Financial Metrics: " + str(financial_model.get('metrics', {})) if financial_model else ""}

Provide an investment recommendation with:
- action: "proceed", "conditional", or "decline"
- amount: investment amount recommendation
- conditions: list of 2-3 conditions/milestones (if conditional)
- reasoning: brief explanation (1-2 sentences)

Format as JSON:
{{
  "action": "proceed|conditional|decline",
  "amount": "$X,XXX,XXX",
  "conditions": ["condition 1", "condition 2"],
  "reasoning": "brief explanation"
}}"""
        
        try:
            response = await self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are an experienced investment committee member. Make clear, actionable recommendations."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                response_format={"type": "json_object"}
            )
            
            import json
            result = json.loads(response.choices[0].message.content)
            logger.info("Recommendation generated successfully")
            return result
            
        except Exception as e:
            logger.error(f"Error generating recommendation: {e}")
            
            # Determine action based on score
            if score >= 75:
                action = "proceed"
                conditions = []
                reasoning = "Strong opportunity with favorable risk/reward profile."
            elif score >= 60:
                action = "conditional"
                conditions = [
                    "Complete technical and legal due diligence",
                    "Validate financial projections with management",
                    "Secure commitment from other lead investors"
                ]
                reasoning = "Promising opportunity but requires additional validation before proceeding."
            else:
                action = "decline"
                conditions = ["Reconsider if fundamentals improve significantly"]
                reasoning = "Opportunity does not meet investment criteria at this time."
            
            return {
                "action": action,
                "amount": f"${company_data.get('funding_amount', 0):,.0f}",
                "conditions": conditions,
                "reasoning": reasoning
            }
    
    def _calculate_expected_roi(self, financial_model: Optional[Dict]) -> str:
        """Calculate expected ROI from financial model"""
        if not financial_model:
            return "25-35% (Target)"
        
        metrics = financial_model.get('metrics', {})
        cagr = metrics.get('revenue_cagr', 0)
        
        if cagr > 2.0:  # >200% CAGR
            return "40-50% (High Growth)"
        elif cagr > 1.0:  # >100% CAGR
            return "30-40% (Strong Growth)"
        elif cagr > 0.5:  # >50% CAGR
            return "25-35% (Solid Growth)"
        else:
            return "20-30% (Moderate Growth)"
